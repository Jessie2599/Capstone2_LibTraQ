import csv
import io
import os
import shutil
import subprocess
import tkinter as tk
from datetime import datetime
from docx2pdf import convert
from tkinter import filedialog
from tkinter import messagebox, ttk
from pathlib import Path
import matplotlib.pyplot as plt
import pymysql
import qrcode
from PIL import Image, ImageTk
from PIL import ImageFile
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.enum.section import WD_ORIENT
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import paragraph
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

ImageFile.LOAD_TRUNCATED_IMAGES = True

PIN = "admin"
my_tree = None
rows = []

global_qr_code_image = None
global_photo_image = None
file_path = None
global background_photo, back_icon


def connection():
    conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
    return conn


def verify_pin(event=None):
    pin = pin_entry.get()
    if pin.strip() == "":  # Check if the pin is an empty string or contains only whitespace
        messagebox.showerror("Error", "Please fill out this field.")
    elif pin == PIN:
        open_next_window()
    else:
        messagebox.showerror("Error", "Your PIN is invalid.")


def open_next_window():
    window.destroy()
    second_window()


def second_window():
    global my_tree

    def confirm_logout(second_window):
        result = messagebox.askquestion("Logout", "Are you sure you want to logout?")
        if result == "yes":
            if second_window:
                second_window.destroy()
                main_window()

    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, text="", values=array, tag="orow")

        my_tree.tag_configure('orow', background='', font=("Arial", 10))
        my_tree.place(x=10, y=150)

    second_window = tk.Tk()
    second_window.geometry("1366x768")
    second_window.resizable(width=False, height=False)
    second_window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")
    second_window.attributes('-fullscreen', True)

    background_image = tk.PhotoImage(file="images/system_background.png")

    background_label = tk.Label(second_window, image=background_image)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)

    title_label = tk.Label(second_window, text="Library Tracker and Monitoring System Using QR Code", bg="Silver",
                           font=("Arial Rounded MT Bold", 30, "bold"), borderwidth=10, relief="ridge")
    title_label.place(x=7, y=10, relwidth=0.99, height=100)

    title_label = tk.Label(second_window, text="Library Attendance", font=("Rockwell Bold", 20))
    title_label.place(x=15, y=110)

    admin_tab = tk.Label(second_window, text="ADMIN",fg="gray", font=("Rockwell Bold", 28))
    admin_tab.place(x=1195, y=110)

    admin_tab = tk.Label(second_window, text="TAB", fg="gray", font=("Rockwell Bold", 28))
    admin_tab.place(x=1235, y=155)

    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance1")
        results = cursor.fetchall()
        conn.commit()
        conn.close()
        return results

    my_tree = ttk.Treeview(second_window, height=28)
    my_tree['columns'] = (
    "Library ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("Library ID Number", anchor="center", width=160)
    my_tree.column("First Name", anchor="center", width=160)
    my_tree.column("Middle Name", anchor="center", width=160)
    my_tree.column("Last Name", anchor="center", width=160)
    my_tree.column("Course", anchor="center", width=150)
    my_tree.column("Purpose", anchor="center", width=200)
    my_tree.column("Date & Time", anchor="center", width=190)

    my_tree.heading("Library ID Number", text="Library ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")

    refreshTable()



    icon_image = Image.open("images/add_user_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 130
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_add_record = tk.Button(second_window, image=icon_photo, height=new_height, width=new_width,
                                  command=open_add_record_window)
    button_add_record.image = icon_photo
    button_add_record.place(x=1210, y=207)

    icon_image = Image.open("images/list_of_user_icon.png")
    new_width, new_height = 130, 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    list_of_students = tk.Button(second_window, image=icon_photo, command=open_list_of_students_window,
                                 height=new_height, width=new_width)
    list_of_students.image = icon_photo
    list_of_students.place(x=1210, y=297)

    icon_image = Image.open("images/generate_report_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 130
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    button_generate_report = tk.Button(second_window, image=icon_photo, command=open_generate_report_window,
                                       height=new_height, width=new_width)
    button_generate_report.image = icon_photo
    button_generate_report.place(x=1210, y=387)

    icon_image = Image.open("images/library_utilization_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 130
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    library_utilization_window = tk.Button(second_window, image=icon_photo, command=open_library_utilization_window,
                                           height=new_height, width=new_width)
    library_utilization_window.image = icon_photo
    library_utilization_window.place(x=1210, y=477)

    icon_image = Image.open("images/about_us_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 130
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    about_us_window = tk.Button(second_window, image=icon_photo, command=open_about_us_window, height=new_height,
                                width=new_width)
    about_us_window.image = icon_photo
    about_us_window.place(x=1210, y=567)

    icon_image = Image.open("images/logout_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 130
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    logout_button = tk.Button(second_window, image=icon_photo, height=new_height, width=new_width,
                              command=lambda: confirm_logout(second_window))
    logout_button.image = icon_photo
    logout_button.place(x=1210, y=657)

    second_window.mainloop()


def open_add_record_window():
    def generate_qr_code():
        global global_qr_code_image
        qr_data = library_id_no_entry.get()

        if qr_data:
            qr = qrcode.QRCode(version=1, box_size=10, border=4)  # Adjust box_size and border here
            qr.add_data(qr_data)
            qr.make(fit=True)
            qr_image = qr.make_image(fill="black", back_color="white")

            qr_image = qr_image.resize((220, 200))  # Adjust size to match photo size

            qr_byte_io = io.BytesIO()
            qr_image.save(qr_byte_io, format="PNG")

            global_qr_code_image = qr_byte_io.getvalue()

            for widget in qr_photo_box.winfo_children():
                widget.destroy()

            qr_label = tk.Label(qr_photo_box)
            qr_label.img = ImageTk.PhotoImage(Image.open(io.BytesIO(global_qr_code_image)))
            qr_label.config(image=qr_label.img)
            qr_label.pack()

    def open_file_dialog():
        global file_path
        file_path = filedialog.askopenfilename(parent=add_record_window,
                                               filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif")])
        if not file_path:
            return

        image = Image.open(file_path)
        image = image.resize((220, 200), Image.LANCZOS)
        global global_photo_image
        global_photo_image = ImageTk.PhotoImage(image)

        for widget in photo_box.winfo_children():
            widget.destroy()

        label = tk.Label(photo_box, image=global_photo_image)
        label.image = global_photo_image
        label.pack()

        add_record_window.deiconify()

    def generate_pdf(library_id_no_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get, photo_data,
                     qr_image_data, header_image_path):
        desktop_path = Path.home() / "Desktop"  # Path to the desktop
        folder_path = desktop_path / "Library Card"  # Path to the folder on the desktop
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)  # Create the folder if it doesn't exist
        filename = f"{lname_get}_{fname_get}.pdf"  # Modified filename
        file_path = folder_path / filename
        c = canvas.Canvas(str(file_path), pagesize=(letter[0] / 2, letter[1] / 2))

        # First Page
        # Draw a rectangular border to resemble an ID card
        c.rect(20, 20, letter[0] / 2 - 40, letter[1] / 2 - 40)

        header_image_path = "images/header.png"
        header_image = Image.open(header_image_path)
        header_image_width, header_image_height = header_image.size
        c.drawInlineImage(header_image, 36, letter[1] / 2 - 80, width=header_image_width / 3,
                          height=header_image_height / 2)

        # Title for the ID card
        c.setFont("Helvetica-Bold", 20)
        c.drawString(100, letter[1] / 2 - 100, "Library Card")


        # Student information
        c.setFont("Helvetica-Bold", 10)
        c.drawString(40, letter[1] / 2 - 240, "ID No:")
        c.setFont("Helvetica", 10)  # Change the font to bold
        c.drawString(80, letter[1] / 2 - 240, str(library_id_no_get))
        c.setFont("Helvetica-Bold", 10)  # Reset the font to regular
        c.drawString(40, letter[1] / 2 - 260, "Name:")
        c.setFont("Helvetica", 10)  # Change the font to bold
        c.drawString(80, letter[1] / 2 - 260, f"{fname_get} {mname_get} {lname_get}")
        c.setFont("Helvetica-Bold", 10)  # Reset the font to regular
        c.drawString(40, letter[1] / 2 - 280, "Sex:")
        c.setFont("Helvetica", 10)  # Change the font to bold
        c.drawString(80, letter[1] / 2 - 280, str(sex_get))
        c.setFont("Helvetica-Bold", 10)  # Reset the font to regular
        c.drawString(40, letter[1] / 2 - 300, "Course:")
        c.setFont("Helvetica", 10)  # Change the font to bold
        c.drawString(80, letter[1] / 2 - 300, str(course_get))
        c.setFont("Helvetica-Bold", 10)  # Reset the font to regular
        c.drawString(40, letter[1] / 2 - 320, "Status:")
        c.setFont("Helvetica", 10)  # Change the font to bold
        c.drawString(80, letter[1] / 2 - 320, str(status_get))
        c.setFont("Helvetica-Bold", 10)  # Reset the font to regular
        c.drawString(40, letter[1] / 2 - 340, "Signature:_______________")

        # Add the photo (size adjusted to fit) with border
        photo = Image.open(io.BytesIO(photo_data))
        photo = photo.convert("RGB")
        photo_width, photo_height = photo.size
        c.setLineWidth(2)
        c.drawImage(ImageReader(photo), 40, letter[1] / 2 - 220, width=photo_width / 2, height=photo_height / 2)
        c.rect(40, letter[1] / 2 - 220, photo_width / 2, photo_height / 2)

        # Add the QR code (size adjusted to fit) with border
        qr_image = Image.open(io.BytesIO(qr_image_data))
        qr_image = qr_image.convert("RGB")
        qr_width, qr_height = qr_image.size
        c.setLineWidth(2)
        c.drawImage(ImageReader(qr_image), 160, letter[1] / 2 - 220, width=qr_width / 2, height=qr_height / 2)
        c.rect(160, letter[1] / 2 - 220, qr_width / 2, qr_height / 2)

        c.save()
    def display_record(library_id_no_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get, photo_data,
                       qr_image_data):
        def on_generate_pdf(header_image_path=None):
            generate_pdf(
                library_id_no_get,
                fname_get,
                mname_get,
                lname_get,
                sex_get,
                course_get,
                status_get,
                photo_data,
                qr_image_data,
                header_image_path,
            )
            file_path = str(Path.home() / "Desktop" / "Library Card" / f"{lname_get}_{fname_get}.pdf")
            subprocess.Popen([file_path], shell=True)
            record_window.destroy()

        record_window = tk.Toplevel()
        record_window.title("Record Details")

        record_info = [
            f"Library ID No.: {library_id_no_get}",
            f"First Name: {fname_get}",
            f"Middle Name: {mname_get}",
            f"Last Name: {lname_get}",
            f"Sex: {sex_get}",
            f"Course: {course_get}",
            f"Status: {status_get}"
        ]

        for info in record_info:
            label = tk.Label(record_window, text=info)
            label.pack()

        photo = Image.open(io.BytesIO(photo_data))
        photo = ImageTk.PhotoImage(photo)
        photo_label = tk.Label(record_window, image=photo)
        photo_label.image = photo
        photo_label.pack()

        qr_image = Image.open(io.BytesIO(qr_image_data))
        qr_image = ImageTk.PhotoImage(qr_image)
        qr_label = tk.Label(record_window, image=qr_image)
        qr_label.image = qr_image
        qr_label.pack()

        pdf_button = tk.Button(record_window, text="Generate PDF", command=on_generate_pdf)
        pdf_button.pack()

    # noinspection PyTypeChecker,PyUnreachableCode
    def add_record_window_save_to_db():
        global global_qr_code_image, global_photo_image

        library_id_no_get = library_id_no_entry.get()
        fname_get = first_name_entry.get()
        mname_get = middle_name_entry.get()
        lname_get = last_name_entry.get()
        sex_get = sex_var.get()
        course_get = course_var.get()
        status_get = status_var.get()

        if not global_qr_code_image:
            messagebox.showerror("Error", "Please generate a QR code.", parent=add_record_window)
            return

        if not global_photo_image:
            messagebox.showerror("Error", "Please upload a photo.", parent=add_record_window)
            return

        if not library_id_no_get or not fname_get or not mname_get or not lname_get:
            messagebox.showerror("Error", "Please fill in all the required fields.")
            return

        try:
            conn = pymysql.connect(host='localhost', user='root', database='libtraq_db')
            cursor = conn.cursor()
            datetime.now().strftime("%Y%m%d%H%M%S")
            qr_image_data = global_qr_code_image
            photo_pil = Image.open(file_path)
            photo_pil = photo_pil.resize((220, 200), Image.LANCZOS)
            photo_buffer = io.BytesIO()
            photo_pil.save(photo_buffer, format="PNG")
            photo_data = photo_buffer.getvalue()

            cursor.execute("INSERT INTO student (library_id_no, first_name, middle_name, last_name, sex, course, status, qr_code, photo) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)", (library_id_no_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get, qr_image_data, photo_data))
            conn.commit()
            conn.close()

            library_id_no_entry.delete(0, tk.END)
            first_name_entry.delete(0, tk.END)
            middle_name_entry.delete(0, tk.END)
            last_name_entry.delete(0, tk.END)
            sex_var.set("Male")
            course_var.set("BSA")
            status_var.set("Active")

            for widget in photo_box.winfo_children():
                widget.destroy()

            for widget in qr_photo_box.winfo_children():
                widget.destroy()

            display_record(library_id_no_get, fname_get, mname_get, lname_get, sex_get, course_get, status_get, photo_data, qr_image_data)

        except Exception as error:
            message = f"Error: {str(error)}"
            messagebox.showerror("Error", message)
            print(error)

    def go_back():
        add_record_window.destroy()

    add_record_window = tk.Toplevel()
    add_record_window.title("Add Users")
    add_record_window.geometry("1368x768")
    add_record_window.attributes("-fullscreen", True)

    app_title_label = tk.Label(add_record_window, text="Library Tracker and Monitoring System Using QR Code",
                               font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(add_record_window, text="Add Student Record", bg="gray", font=("Rockwell", 24, "bold"),
                           borderwidth=1, relief="solid")
    title_label.place(x=7, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(add_record_window, image=back_icon, command=go_back)
    back_button.place(x=10, y=123)

    library_id_no_label = tk.Label(add_record_window, text="Library ID Number:", font=("Rockwell", 18))
    library_id_no_label.place(x=520, y=200)

    library_id_no_entry = tk.Entry(add_record_window, width=45, bg="lightgray", font=("Arial", 18))
    library_id_no_entry.place(x=745, y=200)

    first_name_label = tk.Label(add_record_window, text="First Name:", font=("Rockwell", 18))
    first_name_label.place(x=520, y=240)

    first_name_entry = tk.Entry(add_record_window, width=45, bg="lightgray", font=("Arial", 18))
    first_name_entry.place(x=745, y=240)

    middle_name_label = tk.Label(add_record_window, text="Middle Name:", font=("Rockwell", 18))
    middle_name_label.place(x=520, y=280)

    middle_name_entry = tk.Entry(add_record_window, width=45, bg="lightgray", font=("Arial", 18))
    middle_name_entry.place(x=745, y=280)

    last_name_label = tk.Label(add_record_window, text="Last Name:", font=("Rockwell", 18))
    last_name_label.place(x=520, y=320)

    last_name_entry = tk.Entry(add_record_window, width=45, bg="lightgray", font=("Arial", 18))
    last_name_entry.place(x=745, y=320)

    sex_label = tk.Label(add_record_window, text="Sex:", font=("Arial", 18))
    sex_label.place(x=520, y=360)

    sex_var = tk.StringVar(add_record_window)
    sex_var.set("Male")

    male_radio = tk.Radiobutton(add_record_window, text="Male", variable=sex_var, value="Male", font=("Rockwell", 18))
    female_radio = tk.Radiobutton(add_record_window, text="Female", variable=sex_var, value="Female",
                                  font=("Rockwell", 18))

    male_radio.place(x=670, y=360)
    female_radio.place(x=770, y=360)

    course_label = tk.Label(add_record_window, text="Course:", font=("Rockwell", 18))
    course_label.place(x=520, y=400)

    course_var = tk.StringVar(add_record_window)
    course_var.set("BSA")

    bsa_radio = tk.Radiobutton(add_record_window, text="BSA", variable=course_var, value="BSA", font=("Rockwell", 18))
    bsed_radio = tk.Radiobutton(add_record_window, text="BSED", variable=course_var, value="BSED",
                                font=("Rockwell", 18))
    beed_radio = tk.Radiobutton(add_record_window, text="BEED", variable=course_var, value="BEED",
                                font=("Rockwell", 18))
    bshm_radio = tk.Radiobutton(add_record_window, text="BSHM", variable=course_var, value="BSHM",
                                font=("Rockwell", 18))
    bsoa_radio = tk.Radiobutton(add_record_window, text="BSOA", variable=course_var, value="BSOA",
                                font=("Rockwell", 18))
    bsit_radio = tk.Radiobutton(add_record_window, text="BSIT", variable=course_var, value="BSIT",
                                font=("Rockwell", 18))

    bsa_radio.place(x=670, y=400)
    bsed_radio.place(x=770, y=400)
    beed_radio.place(x=870, y=400)
    bshm_radio.place(x=970, y=400)
    bsoa_radio.place(x=1070, y=400)
    bsit_radio.place(x=1170, y=400)

    status_label = tk.Label(add_record_window, text="Status:", font=("Rockwell", 18))
    status_label.place(x=520, y=440)

    status_var = tk.StringVar(add_record_window)
    status_var.set("Active")

    active_radio = tk.Radiobutton(add_record_window, text="Active", variable=status_var, value="Active",
                                  font=("Rockwell", 18))
    active_radio.place(x=670, y=440)

    photo_box = tk.Canvas(add_record_window, width=220, height=200, bg="white", highlightbackground="black",
                          highlightthickness=2)
    photo_box.place(x=10, y=200)

    qr_photo_box = tk.Canvas(add_record_window, width=220, height=200, bg="white", highlightbackground="black",
                             highlightthickness=2)
    qr_photo_box.place(x=260, y=200)

    upload_button = tk.Button(add_record_window, text="Upload Photo", font=("Arial", 18), command=open_file_dialog)
    upload_button.place(x=38, y=406)

    qr_code_button = tk.Button(add_record_window, text="Generate QR", font=("Arial", 18), command=generate_qr_code)
    qr_code_button.place(x=290, y=406)

    save_button = tk.Button(add_record_window, text="Save Record", bg="gray", font=("Rockwell", 18),
                            command=add_record_window_save_to_db)
    save_button.place(x=670, y=510)

    add_record_window.mainloop()


def open_list_of_students_window():
    # noinspection PyGlobalUndefined
    global table

    def update_student(event):
        # noinspection PyGlobalUndefined
        global table
        selected_item = table.focus()
        if not selected_item:
            return
        selected_student_values = table.item(selected_item, 'values')

        search_entry.delete(0, tk.END)
        search_entry.insert(0, selected_student_values[1])
        selected_student_values[7]

        update_window = tk.Toplevel()
        update_window.title("LibtraQ! Update Student")
        update_window.geometry("800x550")
        update_window.resizable(width=False, height=False)

        library_id_no_label = tk.Label(update_window, text="Library ID Number:", font=("Rockwell", 12))
        library_id_no_label.place(x=330, y=19)
        library_id_no_entry = tk.Entry(update_window, width=35, bg="lightgray", font=("Arial", 12))
        library_id_no_entry.insert(0, selected_student_values[0])
        library_id_no_entry.place(x=240, y=45)

        first_name_label = tk.Label(update_window, text="First Name:", font=("Rockwell", 12))
        first_name_label.place(x=330, y=72)
        first_name_entry = tk.Entry(update_window, width=35, bg="lightgray", font=("Arial", 12))
        first_name_entry.insert(0, selected_student_values[1])
        first_name_entry.place(x=240, y=100)

        middle_name_label = tk.Label(update_window, text="Middle Name:", font=("Rockwell", 12))
        middle_name_label.place(x=330, y=125)
        middle_name_entry = tk.Entry(update_window, width=35, bg="lightgray", font=("Arial", 12))
        middle_name_entry.insert(0, selected_student_values[2])
        middle_name_entry.place(x=240, y=155)

        last_name_label = tk.Label(update_window, text="Last Name:", font=("Rockwell", 12))
        last_name_label.place(x=330, y=178)
        last_name_entry = tk.Entry(update_window, width=35, bg="lightgray", font=("Arial", 12))
        last_name_entry.insert(0, selected_student_values[3])
        last_name_entry.place(x=240, y=205)

        sex_label = tk.Label(update_window, text="Sex:", font=("Rockwell", 12))
        sex_label.place(x=240, y=235)
        sex_value = tk.StringVar()
        sex_value.set(selected_student_values[4])
        male_radio = tk.Radiobutton(update_window, text="Male", variable=sex_value, value="Male")
        female_radio = tk.Radiobutton(update_window, text="Female", variable=sex_value, value="Female")
        male_radio.place(x=300, y=235)
        female_radio.place(x=360, y=235)

        course_label = tk.Label(update_window, text="Course:", font=("Rockwell", 12))
        course_label.place(x=240, y=265)
        course_value = tk.StringVar()
        course_value.set(selected_student_values[5])
        course_options = ["BSA", "BSED", "BEED", "BSHM", "BSOA", "BSIT"]
        x_course = 300
        y_course = 265
        for option in course_options:
            tk.Radiobutton(update_window, text=option, variable=course_value, value=option).place(x=x_course,
                                                                                                  y=y_course)
            y_course += 30  # Adjust the spacing between the radio buttons

        status_label = tk.Label(update_window, text="Status:", font=("Rockwell", 12))
        status_label.place(x=240, y=445)
        status_value = tk.StringVar()
        status_value.set(selected_student_values[6])
        status_options = ["Active", "Inactive"]
        x_status = 300
        y_status = 445
        for option in status_options:
            tk.Radiobutton(update_window, text=option, variable=status_value, value=option).place(x=x_status,
                                                                                                  y=y_status)
            x_status += 60
        update_button = tk.Button(update_window, bg="light Gray", text="Save Update", font=("Rockwell Bold", 14),
                                  command=lambda: perform_update(list_of_students_window, library_id_no_entry,
                                                                 first_name_entry,
                                                                 middle_name_entry, last_name_entry, sex_value,
                                                                 course_value, status_value, update_window))
        update_button.place(x=330, y=490)

        photo_box = tk.Frame(update_window, width=200, height=200, bg="white", relief="solid")
        photo_box.place(x=30, y=20)

        qr_code_box = tk.Frame(update_window, bg="white", relief="solid", width=200, height=200)
        qr_code_box.place(x=570, y=20)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT photo FROM student WHERE library_id_no = %s", (selected_student_values[0],))
        photo_data = cursor.fetchone()
        conn.close()

        if photo_data:
            photo_data = photo_data[0]
            photo = Image.open(io.BytesIO(photo_data))
            photo = photo.resize((200, 200), Image.LANCZOS)
            photo = ImageTk.PhotoImage(photo)

            photo_label = tk.Label(photo_box, image=photo)
            photo_label.photo = photo
            photo_label.pack(fill="both", expand=True)

            conn = connection()
            cursor = conn.cursor()
            cursor.execute("SELECT qr_code FROM student WHERE library_id_no = %s", (selected_student_values[0],))
            qr_data = cursor.fetchone()
            conn.close()

            if qr_data:
                qr_code_data = qr_data[0]
                qr_code = Image.open(io.BytesIO(qr_code_data))
                qr_code = qr_code.resize((200, 200), Image.LANCZOS)
                qr_code = ImageTk.PhotoImage(qr_code)

                qr_label = tk.Label(qr_code_box, image=qr_code)
                qr_label.qr_code = qr_code
                qr_label.pack(fill="both", expand=True)

    def perform_update(list_of_users_window, library_id_no_entry, first_name_entry, middle_name_entry, last_name_entry,
                       sex_value, course_value, status_entry, update_window):
        updated_library_id_no = library_id_no_entry.get()
        updated_first_name = first_name_entry.get()
        updated_middle_name = middle_name_entry.get()
        updated_last_name = last_name_entry.get()
        updated_sex = sex_value.get()
        updated_course = course_value.get()
        updated_status = status_entry.get()

        conn = connection()
        cursor = conn.cursor()

        update_query = "UPDATE student SET first_name = %s, middle_name = %s, last_name = %s, sex = %s, course = %s, status = %s WHERE library_id_no = %s"
        cursor.execute(update_query, (
            updated_first_name, updated_middle_name, updated_last_name, updated_sex, updated_course, updated_status,
            updated_library_id_no))
        conn.commit()
        conn.close()

        # Assuming `table` is a global variable representing your treeview
        populate_treeview(table)

        # Show messagebox in the list_of_users_window
        messagebox.showinfo("Update Successful", "Update Successfully!", parent=list_of_users_window)

        # Destroy the update_window
        update_window.destroy()

    def populate_treeview(table):
        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()
        conn.commit()
        conn.close()

        for row in rows:
            table.insert("", "end", values=row)

    def search_data():
        search_query = search_entry.get().strip().lower()

        for row in table.get_children():
            table.delete(row)

        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM student")
        rows = cursor.fetchall()

        for row in rows:
            if any(search_query in str(cell).strip().lower() for cell in row):
                table.insert("", "end", values=row)

        conn.close()

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    def go_back():
        list_of_students_window.destroy()

    list_of_students_window = tk.Toplevel()
    list_of_students_window.title("List Of Students")
    list_of_students_window.geometry("1368x768")
    list_of_students_window.attributes('-fullscreen', True)

    app_title_label = tk.Label(list_of_students_window, text="Library Tracker and Monitoring System Using QR Code",
                               font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(list_of_students_window, text="List of Student", bg="gray", font=("Rockwell", 24, "bold"),
                           borderwidth=1, relief="solid")
    title_label.place(x=7, y=102, relwidth=0.99, height=80)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(list_of_students_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(list_of_students_window, image=search_photo, command=search_data)
    search_button.place(x=1270, y=190)

    search_entry = tk.Entry(list_of_students_window, font=("Arial", 25), width=69, bg="Light Grey")
    search_entry.place(x=10, y=205)

    table_frame = ttk.Frame(list_of_students_window)
    table_frame.place(x=5, y=270)

    table_columns = ("Library ID Number", "First Name", "Middle Name", "Last Name", "Sex", "Course", "Status")
    table = ttk.Treeview(table_frame, columns=table_columns, show="headings", height=23)
    table.place(x=1000, y=1000)

    scrollbar = ttk.Scrollbar(list_of_students_window, orient="vertical", command=table.yview)
    table.configure(yscrollcommand=scrollbar.set)
    table.pack(fill="both", expand=False)
    scrollbar.pack(side="right", fill="y")

    table.bind("<Double-1>", update_student)

    for col in table_columns:
        table.heading(col, text=col)
        table.column(col, width=100)

    table.heading("Library ID Number", text="Library ID Number")
    table.heading("First Name", text="First Name")
    table.heading("Middle Name", text="Middle Name")
    table.heading("Last Name", text="Last Name")
    table.heading("Sex", text="Sex")
    table.heading("Course", text="Course")
    table.heading("Status", text="Status")

    table.column("Library ID Number", width=200, anchor="center")
    table.column("First Name", width=200, anchor="center")
    table.column("Middle Name", width=200, anchor="center")
    table.column("Last Name", width=200, anchor="center")
    table.column("Sex", width=180, anchor="center")
    table.column("Course", width=180, anchor="center")
    table.column("Status", width=180, anchor="center")

    populate_treeview(table)

    list_of_students_window.mainloop()


def open_generate_report_window():
    generate_report_window = tk.Toplevel()
    generate_report_window.title("Generate Report")
    generate_report_window.geometry("1368x768")
    generate_report_window.attributes('-fullscreen', True)

    def refreshTable():
        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in read():
            my_tree.insert(parent='', index='end', iid=array, values=array, tag='orow')
            my_tree.place(x=10, y=300)

    def read():
        conn = connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM library_attendance1")
        results = cursor.fetchall()
        conn.close()
        return results

    def connection():
        conn = pymysql.connect(host="localhost", user="root", password="", database="libtraq_db")
        return conn

    # noinspection PyGlobalUndefined
    def filter_data():
        global results
        selected_course = course_var.get()
        selected_semester = semester_var.get()
        selected_academic_year = academic_year_var.get()

        conn = connection()
        cursor = conn.cursor()

        try:
            if selected_course == "All Courses" and selected_semester == "Whole School Year" and selected_academic_year == "All Academic Year":
                cursor.execute("SELECT * FROM library_attendance1")

            elif selected_course == "All Courses" and selected_academic_year == "All Academic Year":
                cursor.execute("SELECT * FROM library_attendance1 WHERE Semester=%s", (selected_semester,))

            elif selected_semester == "Whole School Year" and selected_academic_year == "All Academic Year":
                cursor.execute("SELECT * FROM library_attendance1 WHERE Course=%s", (selected_course,))

            elif selected_course == "All Courses" and selected_semester == "Whole School Year":
                cursor.execute("SELECT * FROM library_attendance1 WHERE Academic_Year=%s", (selected_academic_year,))

            elif selected_academic_year == "All Academic Year":
                cursor.execute("SELECT * FROM library_attendance1 WHERE Course=%s AND Semester=%s", (selected_course, selected_semester))

            elif selected_course == "All Courses":
                cursor.execute("SELECT * FROM library_attendance1 WHERE Semester=%s AND Academic_Year=%s", (selected_semester, selected_academic_year))

            else:
                cursor.execute("SELECT * FROM library_attendance1 WHERE Course=%s AND Semester=%s AND Academic_Year=%s", (selected_course, selected_semester, selected_academic_year,))

            results = cursor.fetchall()
            conn.close()

        except Exception as e:
            print(f"Error occurred: {e}")

        for date in my_tree.get_children():
            my_tree.delete(date)

        for array in results:
            my_tree.insert(parent='', index='end', iid=array, values=array, tag='orow')
            my_tree.place(x=10, y=300)

    def go_back():
        generate_report_window.destroy()

    def print_report():
        course_var.get()

        filename = "Library_Report.docx"

        doc = Document()

        section = doc.sections[0]
        new_width, new_height = section.page_width, section.page_height  # Swap the width and height values
        section.orientation = WD_ORIENT.PORTRAIT  # Set the orientation to portrait
        section.page_width = new_width
        section.page_height = new_height

        header = section.header
        paragraph = header.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('images/header.png', width=Inches(6.48))

        doc.add_heading("Library Utilization Report", 0)
        doc.add_paragraph("Course: " + course_var.get())
        doc.add_paragraph("Semester: " + semester_var.get())
        doc.add_paragraph("Academic Year: " + academic_year_var.get())
        doc.add_paragraph("Generated by: Mariter Asur")

        table = doc.add_table(rows=1, cols=7)
        table.allow_autofit = False
        table.style = 'Table Grid'

        table_headers = table.rows[0].cells
        table_headers[0].text = "Library ID Number"
        table_headers[1].text = "First Name"
        table_headers[2].text = "Middle Name"
        table_headers[3].text = "Last Name"
        table_headers[4].text = "Course"
        table_headers[5].text = "Purpose"
        table_headers[6].text = "Date & Time"

        for item in my_tree.get_children():
            values = my_tree.item(item, 'values')
            data = list(map(str, values))
            table.add_row().cells[0].text = data[0]
            table.rows[-1].cells[1].text = data[1]
            table.rows[-1].cells[2].text = data[2]
            table.rows[-1].cells[3].text = data[3]
            table.rows[-1].cells[4].text = data[4]
            table.rows[-1].cells[5].text = data[5]
            table.rows[-1].cells[6].text = data[6]

        doc.save(filename)

        downloads_folder = os.path.expanduser("~") + "/Desktop/Ready to Print Reports"
        doc_path = os.path.join(downloads_folder, filename)
        shutil.move(filename, doc_path)

        os.startfile(doc_path)

        print(f"Report generated and saved as {doc_path}")

    app_title_label = tk.Label(generate_report_window, text="Library Tracker and Monitoring System Using QR Code",
                               font=("Arial Rounded MT Bold", 30, "bold"))
    app_title_label.place(x=0, y=0, relwidth=0.99, height=100)

    title_label = tk.Label(generate_report_window, text="Generate Report", bg="gray", font=("Rockwell", 24, "bold"),
                           borderwidth=1, relief="solid")
    title_label.place(x=7, y=102, relwidth=0.99, height=80)

    semester_label = tk.Label(generate_report_window, text="Semester:", font=("Rockwell", 18))
    semester_label.place(x=10, y=210)

    course_label = tk.Label(generate_report_window, text="Courses:", font=("Rockwell", 18))
    course_label.place(x=480, y=210)

    semester_var = tk.StringVar(generate_report_window)
    semester_choices = ["1st Semester", "2nd Semester", "Whole School Year"]
    semester_dropdown = ttk.Combobox(generate_report_window, textvariable=semester_var, values=semester_choices,
                                     font=("Rockwell Bold", 16), width=25)
    semester_dropdown.place(x=130, y=215)

    course_var = tk.StringVar(generate_report_window)
    course_choices = ["BSA", "BSED", "BEED", "BSHM", "BSOA", "BSIT", "All Courses"]
    course_dropdown = ttk.Combobox(generate_report_window, textvariable=course_var, value=course_choices,
                                   font=("Rockwell Bold", 16), width=25)
    course_dropdown.place(x=590, y=215)

    def generate_academic_years(start_year, end_year):
        academic_years = []
        for year in range(start_year, end_year + 1):
            academic_years.append(f"{year}-{year + 1}")
        return academic_years

    academic_year_label = tk.Label(generate_report_window, text="Academic Year:", font=("Rockwell", 18))
    academic_year_label.place(x=10, y=260)

    academic_year_var = tk.StringVar(generate_report_window)
    academic_year_choices = generate_academic_years(2023, 2050)
    academic_year_dropdown = ttk.Combobox(generate_report_window, textvariable=academic_year_var,
                                          values=academic_year_choices, font=("Rockwell Bold", 16), width=25)
    academic_year_dropdown.place(x=200, y=260)

    my_tree = ttk.Treeview(generate_report_window, height=22)
    my_tree.place(x=10, y=260)

    my_tree['columns'] = (
    "Library ID Number", "First Name", "Middle Name", "Last Name", "Course", "Purpose", "Date & Time", "Academic Year",
    "Semester")

    my_tree.column("#0", width=0, stretch=tk.NO)
    my_tree.column("Library ID Number", anchor="center", width=130)
    my_tree.column("First Name", anchor="center", width=150)
    my_tree.column("Middle Name", anchor="center", width=150)
    my_tree.column("Last Name", anchor="center", width=150)
    my_tree.column("Course", anchor="center", width=150)
    my_tree.column("Purpose", anchor="center", width=185)
    my_tree.column("Date & Time", anchor="center", width=150)
    my_tree.column("Academic Year", anchor="center", width=125)
    my_tree.column("Semester", anchor="center", width=150)

    my_tree.heading("Library ID Number", text="ID Number", anchor="center")
    my_tree.heading("First Name", text="First Name", anchor="center")
    my_tree.heading("Middle Name", text="Middle Name", anchor="center")
    my_tree.heading("Last Name", text="Last Name", anchor="center")
    my_tree.heading("Course", text="Course", anchor="center")
    my_tree.heading("Purpose", text="Purpose", anchor="center")
    my_tree.heading("Date & Time", text="Date & Time", anchor="center")
    my_tree.heading("Academic Year", text="Academic Year", anchor="center")
    my_tree.heading("Semester", text="Semester", anchor="center")

    my_tree.place()

    refreshTable()

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(generate_report_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=123)

    search_image = Image.open("images/search_icon.png")
    search_photo = ImageTk.PhotoImage(search_image)

    search_button = tk.Button(generate_report_window, image=search_photo, command=filter_data)
    search_button.place(x=1245, y=190)

    pr = tk.Button(generate_report_window, text="Download Report", command=print_report, font=("Rockwell Bold", 18),
                   bg="Gray")
    pr.place(x=1000, y=200)

    generate_report_window.mainloop()
def open_library_utilization_window():
    library_utilization_window = tk.Toplevel()
    library_utilization_window.title("Library Utilization")
    library_utilization_window.geometry("1366x768")
    library_utilization_window.attributes('-fullscreen', True)

    def go_back():
        library_utilization_window.destroy()

    def fetch_data_and_create_graph():
        try:
            connect = pymysql.connect(host='localhost', user='root', password="", database='libtraq_db')
            cursor = connect.cursor()

            # Graph 1: Attendance by Course
            cursor.execute("SELECT course, COUNT(*) FROM library_attendance1 GROUP BY course")
            results_course = cursor.fetchall()

            # Graph 2: Attendance by Purpose
            cursor.execute("SELECT purpose, COUNT(*) FROM library_attendance1 GROUP BY purpose")
            results_purpose = cursor.fetchall()

            connect.close()

            if not results_course or not results_purpose:
                messagebox.showerror("Error", "No attendance data found.")
                return

            # Processing for Graph 1
            courses, counts_course = zip(*results_course)
            total_count_course = sum(counts_course)
            percentages_course = [count / total_count_course * 100 for count in counts_course]

            # Processing for Graph 2
            purposes, counts_purpose = zip(*results_purpose)
            total_count_purpose = sum(counts_purpose)
            percentages_purpose = [count / total_count_purpose * 100 for count in counts_purpose]

            # Plotting Graph 1
            fig1, ax1 = plt.subplots(figsize=(6, 4))
            bars1 = ax1.bar(courses, counts_course, color='Blue')
            ax1.set_xlabel('Course', fontsize=10)
            ax1.set_ylabel('Attendance Count', fontsize=10)
            ax1.set_title('Attendance by Course', fontsize=12)
            ax1.tick_params(axis='x', rotation=45, labelsize=8)

            for bar, percentage in zip(bars1, percentages_course):
                height = bar.get_height()
                ax1.annotate(f'{percentage:.2f}%', (bar.get_x() + bar.get_width() / 2, height), ha='center', va='bottom', fontsize=8)

            canvas1 = FigureCanvasTkAgg(fig1, master=library_utilization_window)
            canvas1.get_tk_widget().place(x=10, y=350)
            canvas1.draw()

            fig2, ax2 = plt.subplots(figsize=(8, 4))
            ax2.set_title('Attendance by Purpose', fontsize=12)

            # Creating a pie chart
            ax2.pie(counts_purpose, labels=purposes, autopct='%1.1f%%',
                    colors=['red', 'green', 'blue', 'yellow', 'orange'], startangle=90)

            canvas2 = FigureCanvasTkAgg(fig2, master=library_utilization_window)
            canvas2.get_tk_widget().place(x=560, y=350)
            canvas2.draw()

        except pymysql.Error as e:
            messagebox.showerror("Database Error", f"Error: {e}")

    def generate_leaderboard():
        try:
            connect = pymysql.connect(host='localhost', user='root', password="", database='libtraq_db')
            cursor = connect.cursor()

            cursor.execute(
                "SELECT first_name, middle_name, last_name, course, purpose, COUNT(*) as attendance_count FROM library_attendance1 GROUP BY first_name, middle_name, last_name, course, purpose ORDER BY attendance_count DESC")
            results = cursor.fetchall()
            connect.close()

            if not results:
                messagebox.showerror("Error", "No attendance data found.")
                return


            leaderboard_tree = ttk.Treeview(library_utilization_window,
                                            columns=("Full Name", "Course", "Purpose", "Attendance Count"),
                                            show="headings", height=12)
            leaderboard_tree.heading("Full Name", text="Full Name", anchor="center")
            leaderboard_tree.heading("Course", text="Course", anchor="center")
            leaderboard_tree.heading("Purpose", text="Purpose", anchor="center")
            leaderboard_tree.heading("Attendance Count", text="Attendance Count", anchor="center")

            style = ttk.Style()
            style.configure("Treeview.Heading")

            leaderboard_tree.column("Full Name", width=240, anchor='center')
            leaderboard_tree.column("Course", width=170, anchor='center')
            leaderboard_tree.column("Purpose", width=190, anchor='center')
            leaderboard_tree.column("Attendance Count", width=160, anchor='center')
            leaderboard_tree.place(x=300, y=70)

            for index, (first_name, middle_name, last_name, course, purpose, attendance_count) in enumerate(results,
                                                                                                            start=1):
                name = f"{first_name} {middle_name} {last_name}"
                leaderboard_tree.insert("", "end", values=(name, course, purpose, attendance_count))

        except pymysql.Error as e:
            messagebox.showerror("Database Error", f"Error: {e}")

    fetch_data_and_create_graph()
    generate_leaderboard()

    icon_image = Image.open("images/leaderboard_icon.png")
    original_width, original_height = icon_image.size
    original_width / original_height
    new_width = 150
    new_height = 70
    icon_image = icon_image.resize((new_width, new_height), Image.LANCZOS)
    icon_photo = ImageTk.PhotoImage(icon_image)

    refresh = tk.Button(library_utilization_window, image=icon_photo,
                        command=lambda: [fetch_data_and_create_graph(), generate_leaderboard()], height=new_height,
                        width=new_width)
    refresh.image = icon_photo
    refresh.place(x=25, y=70)

    title_label = tk.Label(library_utilization_window, text="Library Utilization",font=("Rockwell", 24, "bold"))
    title_label.place(x=0, y=0, relwidth=0.99, height=60)

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(library_utilization_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=10, y=10)

    library_utilization_window.mainloop()


def open_about_us_window():
    global background_photo, back_icon
    about_us_window = tk.Toplevel()
    about_us_window.title("About Us")
    about_us_window.geometry("1366x768")
    about_us_window.attributes('-fullscreen', True)

    def go_back():
        about_us_window.destroy()

    background_image = Image.open("images/about_us_background.png")
    background_image = background_image.resize((1366, 768), Image.LANCZOS)
    background_photo = ImageTk.PhotoImage(background_image)

    background_label = tk.Label(about_us_window, image=background_photo)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)
    about_us_window.background = background_photo

    back_icon = ImageTk.PhotoImage(Image.open("images/back_icon.png"))
    back_button = tk.Button(about_us_window, image=back_icon, command=go_back, bd=0)
    back_button.place(x=15, y=20)

    about_us_window.mainloop()


# noinspection PyGlobalUndefined
def main_window():
    def on_entry_click(event):
        if pin_entry.get() == hint_text:
            pin_entry.delete(0, tk.END)
            pin_entry.config(fg='black')

    def on_focus_out(event):
        if not pin_entry.get():
            pin_entry.insert(0, hint_text)
            pin_entry.config(fg='grey')

    global window
    window = tk.Tk()
    window.title("LibTraQ: Library Tracker and Monitoring System using QR Code")
    window.geometry("1366x768")
    # window.attributes('-fullscreen', True)

    hint_text = "Enter PIN"

    background_image = tk.PhotoImage(file="images/admin_background.png")

    background_label = tk.Label(window, image=background_image)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)

    pin_label = tk.Label(window, text="Enter PIN", font=("Rockwell", 26), justify='center', bg="#ECECEC")
    pin_label.pack(pady=(540, 10))

    global pin_entry
    pin_entry = tk.Entry(window, show="*", font=("Arial", 30), fg='grey', justify='center')
    pin_entry.insert(0, hint_text)
    pin_entry.bind("<FocusIn>", on_entry_click)
    pin_entry.bind("<FocusOut>", on_focus_out)
    pin_entry.pack(pady=10)
    pin_entry.bind("<Return>", verify_pin)
    pin_entry.focus()

    login_button = tk.Button(window, text="Login", font=("Rockwell", 24), command=verify_pin)
    login_button.pack(pady=0)

    window.mainloop()


if __name__ == '__main__':
    main_window()